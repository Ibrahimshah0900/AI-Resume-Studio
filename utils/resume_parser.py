
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Optional

import pymupdf
from docx import Document

from utils.ocr_parser import extract_image_text


DEFAULT_MIN_TEXT_LENGTH = 50

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
}


@dataclass
class ResumeExtractionResult:
    text: str
    page_count: int
    pages_with_text: int
    character_count: int
    extraction_method: str
    success: bool
    ocr_pages: int = 0
    file_type: str = ""
    error: Optional[str] = None

    @property
    def has_text(self) -> bool:
        return bool(self.text.strip())

    @property
    def used_ocr(self) -> bool:
        return self.ocr_pages > 0


def _failure_result(error: str, file_type: str = "") -> ResumeExtractionResult:
    return ResumeExtractionResult(
        text="",
        page_count=0,
        pages_with_text=0,
        character_count=0,
        extraction_method="none",
        success=False,
        file_type=file_type,
        error=error,
    )


def extract_pdf_text(file_path: str | Path) -> ResumeExtractionResult:
    path = Path(file_path)

    if not path.exists():
        return _failure_result(f"File does not exist: {path}", "pdf")
    if not path.is_file():
        return _failure_result(f"Path is not a file: {path}", "pdf")
    if path.suffix.lower() != ".pdf":
        return _failure_result("The supplied file is not a PDF.", "pdf")

    try:
        document = pymupdf.open(path)
        page_texts: list[str] = []
        pages_with_text = 0

        for page in document:
            text = page.get_text("text").strip()
            if text:
                pages_with_text += 1
            page_texts.append(text)

        page_count = len(document)
        document.close()

        combined_text = "\n\n".join(text for text in page_texts if text).strip()

        return ResumeExtractionResult(
            text=combined_text,
            page_count=page_count,
            pages_with_text=pages_with_text,
            character_count=len(combined_text),
            extraction_method="direct_pdf",
            success=True,
            file_type="pdf",
        )

    except Exception as exc:
        return _failure_result(f"PDF extraction failed: {exc}", "pdf")


def _render_pdf_page(page: pymupdf.Page, output_path: Path, dpi: int = 200) -> None:
    zoom = dpi / 72
    matrix = pymupdf.Matrix(zoom, zoom)
    pixmap = page.get_pixmap(matrix=matrix, alpha=False)
    pixmap.save(str(output_path))


def extract_pdf_with_ocr_fallback(
    file_path: str | Path,
    min_text_length: int = DEFAULT_MIN_TEXT_LENGTH,
    dpi: int = 200,
) -> ResumeExtractionResult:
    path = Path(file_path)
    direct_result = extract_pdf_text(path)

    if not direct_result.success:
        return direct_result

    if direct_result.character_count >= min_text_length:
        return direct_result

    try:
        document = pymupdf.open(path)
        ocr_texts: list[str] = []
        ocr_pages = 0

        temp_directory = path.parent / ".ocr_temp"
        temp_directory.mkdir(parents=True, exist_ok=True)

        for page_index, page in enumerate(document):
            image_path = temp_directory / f"{path.stem}_page_{page_index + 1}.png"
            _render_pdf_page(page, image_path, dpi=dpi)

            ocr_result = extract_image_text(image_path)

            if ocr_result.success and ocr_result.text.strip():
                ocr_texts.append(ocr_result.text.strip())
                ocr_pages += 1

            try:
                image_path.unlink()
            except OSError:
                pass

        document.close()

        try:
            temp_directory.rmdir()
        except OSError:
            pass

        combined_ocr_text = "\n\n".join(ocr_texts).strip()

        if not combined_ocr_text:
            return _failure_result(
                "No text could be extracted using direct extraction or OCR.",
                "pdf",
            )

        return ResumeExtractionResult(
            text=combined_ocr_text,
            page_count=direct_result.page_count,
            pages_with_text=direct_result.pages_with_text,
            character_count=len(combined_ocr_text),
            extraction_method="ocr",
            success=True,
            ocr_pages=ocr_pages,
            file_type="pdf",
        )

    except Exception as exc:
        return _failure_result(f"OCR fallback failed: {exc}", "pdf")


def extract_docx_text(file_path: str | Path) -> ResumeExtractionResult:
    path = Path(file_path)

    if not path.exists():
        return _failure_result(f"File does not exist: {path}", "docx")
    if not path.is_file():
        return _failure_result(f"Path is not a file: {path}", "docx")
    if path.suffix.lower() != ".docx":
        return _failure_result("The supplied file is not a DOCX file.", "docx")

    try:
        document = Document(path)
        text_parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))

        combined_text = "\n".join(text_parts).strip()

        return ResumeExtractionResult(
            text=combined_text,
            page_count=1,
            pages_with_text=1 if combined_text else 0,
            character_count=len(combined_text),
            extraction_method="docx",
            success=True,
            file_type="docx",
        )

    except Exception as exc:
        return _failure_result(f"DOCX extraction failed: {exc}", "docx")


def extract_image_text_file(file_path: str | Path) -> ResumeExtractionResult:
    path = Path(file_path)

    if not path.exists():
        return _failure_result(f"File does not exist: {path}", "image")
    if not path.is_file():
        return _failure_result(f"Path is not a file: {path}", "image")
    if path.suffix.lower() not in {".png", ".jpg", ".jpeg"}:
        return _failure_result("Unsupported image format.", "image")

    try:
        result = extract_image_text(path)

        return ResumeExtractionResult(
            text=result.text,
            page_count=1,
            pages_with_text=1 if result.text.strip() else 0,
            character_count=len(result.text),
            extraction_method="ocr",
            success=result.success,
            ocr_pages=1 if result.success else 0,
            file_type="image",
            error=getattr(result, "error", None),
        )

    except Exception as exc:
        return _failure_result(f"Image OCR failed: {exc}", "image")


def parse_resume_file(file_path: str | Path) -> ResumeExtractionResult:
    path = Path(file_path)

    if not path.exists():
        return _failure_result(f"File does not exist: {path}")

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        return _failure_result(
            f"Unsupported file format: {extension or '[no extension]'}. "
            f"Supported formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if extension == ".pdf":
        return extract_pdf_with_ocr_fallback(path)

    if extension == ".docx":
        return extract_docx_text(path)

    if extension in {".png", ".jpg", ".jpeg"}:
        return extract_image_text_file(path)

    return _failure_result("Unsupported file format.")
