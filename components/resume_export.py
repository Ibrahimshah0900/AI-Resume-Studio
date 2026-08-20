
from __future__ import annotations
import streamlit as st
import uuid
import time
import tempfile
from pathlib import Path
from utils.template_renderer import render_resume_html
from utils.pdf_generator import generate_resume_pdf
from utils.template_engine import get_template, DEFAULT_TEMPLATE_ID

def render_resume_export() -> None:
    st.markdown('<div class="section-title">📥 Export Resume</div>', unsafe_allow_html=True)
    st.write("Generate professional resumes in multiple formats using your selected template.")

    resume_data = st.session_state.resume_data

    if not resume_data:
        st.info("Please create or load a resume in the Resume Builder first.")
        return

    selected_template_id = st.session_state.get("selected_template_id", DEFAULT_TEMPLATE_ID)

    try:
        template = get_template(selected_template_id)
        st.subheader(f"Selected Template: {template.name}")
    except:
        st.subheader(f"Selected Template: {selected_template_id}")
    
    st.write("Your resume will be exported using this template.")
    st.divider()

    # Export options
    export_format = st.radio(
        "Select Export Format",
        ["📑 PDF", "📝 DOCX (Word)"],
        horizontal=True
    )

    if export_format == "📄 HTML":
        st.subheader("Export as HTML")
        if st.button("Generate HTML", key="generate_html", type="primary"):
            try:
                with st.spinner("Generating HTML..."):
                    html_content = render_resume_html(
                        resume_data.model_dump(mode="json"),
                        template_id=selected_template_id,
                    )
                    st.session_state.html_content = html_content
                    st.success("HTML generated successfully!")
            except Exception as e:
                st.error(f"Error generating HTML: {e}")

        if st.session_state.get("html_content"):
            st.download_button(
                label="📥 Download HTML Resume",
                data=st.session_state.html_content,
                file_name="resume.html",
                mime="text/html",
                key=f"download_html_{st.session_state.download_key}",
            )

    elif export_format == "📑 PDF":
        st.subheader("Export as PDF")
        if st.button("Generate PDF", key="generate_pdf", type="primary"):
            try:
                with st.spinner("Generating PDF..."):
                    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                        pdf_path = generate_resume_pdf(
                            resume_data.model_dump(mode="json"),
                            output_path=tmp.name,
                            template_id=selected_template_id,
                        )
                        
                        with open(pdf_path, "rb") as f:
                            pdf_bytes = f.read()
                        
                        st.session_state.pdf_bytes = pdf_bytes
                        st.session_state.pdf_ready = True
                        Path(pdf_path).unlink(missing_ok=True)
                        
                    st.success("PDF generated successfully!")
            except Exception as e:
                st.error(f"Error generating PDF: {e}")

        if st.session_state.get("pdf_ready") and st.session_state.get("pdf_bytes"):
            st.download_button(
                label="📥 Download PDF Resume",
                data=st.session_state.pdf_bytes,
                file_name="resume.pdf",
                mime="application/pdf",
                key=f"download_pdf_{st.session_state.download_key}",
            )

    elif export_format == "📝 DOCX (Word)":
        st.subheader("📝 Export as Word Document")
        st.write("Generate a fully formatted Word document with all your resume data.")
        
        if st.button("📄 Generate DOCX", key="generate_docx", type="primary"):
            try:
                from docx import Document
                from docx.shared import Pt, Inches, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from io import BytesIO
                from utils.data_models import Resume
                
                # Convert Resume object to dict if needed
                if isinstance(resume_data, Resume):
                    resume_dict = resume_data.model_dump()
                else:
                    resume_dict = resume_data
                
                doc = Document()
                
                # Set up document style
                style = doc.styles['Normal']
                style.font.name = 'Calibri'
                style.font.size = Pt(11)
                
                # ----- HEADER: Name -----
                info = resume_dict.get("personal_info", {})
                name = info.get("full_name", info.get("name", "Resume"))
                title = doc.add_heading(name, 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # ----- Professional Title -----
                professional_title = info.get("professional_title", "")
                if professional_title:
                    p = doc.add_paragraph(professional_title)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].font.size = Pt(14)
                    p.runs[0].font.italic = True
                
                # ----- Contact Info -----
                contact_parts = []
                if info.get("email"):
                    contact_parts.append(info.get("email"))
                if info.get("phone"):
                    contact_parts.append(info.get("phone"))
                if info.get("location"):
                    contact_parts.append(info.get("location"))
                if info.get("linkedin"):
                    contact_parts.append(info.get("linkedin"))
                
                if contact_parts:
                    p = doc.add_paragraph(" | ".join(contact_parts))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].font.size = Pt(10)
                
                doc.add_paragraph()
                
                # ----- SUMMARY -----
                summary = resume_dict.get("summary", "")
                if summary:
                    doc.add_heading("Professional Summary", level=1)
                    doc.add_paragraph(summary)
                    doc.add_paragraph()
                
                # ----- EXPERIENCE -----
                experience = resume_dict.get("experience", [])
                if experience:
                    doc.add_heading("Experience", level=1)
                    for exp in experience:
                        # Job title and company
                        job_title = exp.get("job_title", "")
                        company = exp.get("company", "")
                        if job_title and company:
                            p = doc.add_paragraph()
                            run = p.add_run(f"{job_title}")
                            run.bold = True
                            p.add_run(f" — {company}")
                        
                        # Dates
                        start_date = exp.get("start_date", "")
                        end_date = exp.get("end_date", "")
                        if start_date or end_date:
                            p = doc.add_paragraph(f"{start_date} - {end_date}", style='Normal')
                            p.runs[0].font.italic = True
                            p.runs[0].font.size = Pt(10)
                        
                        # Bullet points
                        bullets = exp.get("bullet_points", [])
                        if bullets:
                            for bullet in bullets:
                                if bullet:
                                    doc.add_paragraph(bullet, style='List Bullet')
                        doc.add_paragraph()
                
                # ----- EDUCATION -----
                education = resume_dict.get("education", [])
                if education:
                    doc.add_heading("Education", level=1)
                    for edu in education:
                        degree = edu.get("degree", "")
                        institution = edu.get("institution", "")
                        if degree and institution:
                            p = doc.add_paragraph()
                            run = p.add_run(f"{degree}")
                            run.bold = True
                            p.add_run(f" — {institution}")
                        
                        field = edu.get("field_of_study", "")
                        if field:
                            doc.add_paragraph(field)
                        
                        dates = []
                        if edu.get("start_date"):
                            dates.append(edu.get("start_date"))
                        if edu.get("end_date"):
                            dates.append(edu.get("end_date"))
                        if dates:
                            p = doc.add_paragraph(" - ".join(dates))
                            p.runs[0].font.italic = True
                            p.runs[0].font.size = Pt(10)
                        doc.add_paragraph()
                
                # ----- SKILLS -----
                skills = resume_dict.get("skills", [])
                if skills:
                    doc.add_heading("Skills", level=1)
                    # Add skills as a comma-separated list
                    doc.add_paragraph(", ".join(skills))
                    doc.add_paragraph()
                
                # ----- PROJECTS -----
                projects = resume_dict.get("projects", [])
                if projects:
                    doc.add_heading("Projects", level=1)
                    for proj in projects:
                        proj_name = proj.get("name", "")
                        if proj_name:
                            p = doc.add_paragraph()
                            run = p.add_run(proj_name)
                            run.bold = True
                        
                        if proj.get("description"):
                            doc.add_paragraph(proj.get("description"))
                        
                        techs = proj.get("technologies", [])
                        if techs:
                            doc.add_paragraph(f"Technologies: {', '.join(techs)}")
                        doc.add_paragraph()
                
                # ----- CERTIFICATIONS -----
                certifications = resume_dict.get("certifications", [])
                if certifications:
                    doc.add_heading("Certifications", level=1)
                    for cert in certifications:
                        cert_name = cert.get("name", "")
                        if cert_name:
                            doc.add_paragraph(cert_name)
                    doc.add_paragraph()
                
                # ----- LANGUAGES -----
                languages = resume_dict.get("languages", [])
                if languages:
                    doc.add_heading("Languages", level=1)
                    for lang in languages:
                        lang_name = lang.get("name", "")
                        proficiency = lang.get("proficiency", "")
                        if lang_name:
                            if proficiency:
                                doc.add_paragraph(f"{lang_name} - {proficiency}")
                            else:
                                doc.add_paragraph(lang_name)
                    doc.add_paragraph()
                
                # ----- ACHIEVEMENTS -----
                achievements = resume_dict.get("achievements", [])
                if achievements:
                    doc.add_heading("Achievements", level=1)
                    for ach in achievements:
                        ach_title = ach.get("title", "")
                        if ach_title:
                            doc.add_paragraph(ach_title)
                    doc.add_paragraph()
                
                # Save to bytes
                doc_bytes = BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                
                st.session_state.docx_bytes = doc_bytes.getvalue()
                st.success("✅ DOCX generated successfully with all resume data!")
                
            except ImportError:
                st.error("❌ python-docx is not installed. Please install it with: pip install python-docx")
            except Exception as e:
                st.error(f"❌ Error generating DOCX: {e}")
                import traceback
                st.code(traceback.format_exc())
        
        if st.session_state.get("docx_bytes"):
            st.download_button(
                label="📥 Download DOCX Resume",
                data=st.session_state.docx_bytes,
                file_name="resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_docx_{st.session_state.download_key}",
            )

        if st.session_state.get("docx_bytes"):
            st.download_button(
                label="📥 Download DOCX Resume",
                data=st.session_state.docx_bytes,
                file_name="resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_docx_{st.session_state.download_key}",
            )

    elif export_format == "📋 TXT (ATS)":
        st.subheader("Export as TXT (ATS-Optimized)")
        
        if st.button("Generate TXT", key="generate_txt", type="primary"):
            try:
                # Generate plain text version
                lines = []
                
                info = resume_data.get("personal_info", {})
                lines.append(info.get("full_name", ""))
                lines.append(info.get("professional_title", ""))
                lines.append(f"{info.get('email', '')} | {info.get('phone', '')} | {info.get('location', '')}")
                lines.append("")
                
                if resume_data.get("summary"):
                    lines.append("PROFESSIONAL SUMMARY")
                    lines.append(resume_data.get("summary"))
                    lines.append("")
                
                if resume_data.get("skills"):
                    lines.append("SKILLS")
                    lines.append(", ".join(resume_data.get("skills", [])))
                    lines.append("")
                
                if resume_data.get("experience"):
                    lines.append("EXPERIENCE")
                    for exp in resume_data.get("experience", []):
                        lines.append(f"{exp.get('job_title', '')} - {exp.get('company', '')}")
                        for bullet in exp.get("bullet_points", []):
                            lines.append(f"  • {bullet}")
                        lines.append("")
                
                if resume_data.get("education"):
                    lines.append("EDUCATION")
                    for edu in resume_data.get("education", []):
                        lines.append(f"{edu.get('degree', '')} - {edu.get('institution', '')}")
                        lines.append("")
                
                txt_content = "\n".join(lines)
                st.session_state.txt_content = txt_content
                st.success("TXT generated successfully!")
            except Exception as e:
                st.error(f"Error generating TXT: {e}")

        if st.session_state.get("txt_content"):
            st.download_button(
                label="📥 Download TXT Resume",
                data=st.session_state.txt_content,
                file_name="resume.txt",
                mime="text/plain",
                key=f"download_txt_{st.session_state.download_key}",
            )

    st.divider()
    st.caption("💡 Tip: Select a different template in the Templates section before exporting.")
